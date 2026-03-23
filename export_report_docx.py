from pathlib import Path

from docx import Document


def clean_inline_markdown(text: str) -> str:
    return (
        text.replace("**", "")
        .replace("`", "")
        .replace("  ", " ")
        .strip()
    )


def main() -> None:
    md_path = Path("Звіт_лабораторна_робота_2.md")
    docx_path = Path("Звіт_лабораторна_робота_2.docx")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()

    in_code_block = False
    code_buffer: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                paragraph = document.add_paragraph("\n".join(code_buffer))
                paragraph.style = "No Spacing"
                in_code_block = False
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.strip() == "---":
            document.add_paragraph("")
            continue

        if line.startswith("# "):
            document.add_heading(clean_inline_markdown(line[2:]), level=1)
            continue

        if line.startswith("## "):
            document.add_heading(clean_inline_markdown(line[3:]), level=2)
            continue

        if line.startswith("### "):
            document.add_heading(clean_inline_markdown(line[4:]), level=3)
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            document.add_paragraph(clean_inline_markdown(stripped[2:]), style="List Bullet")
            continue

        if stripped[:2].isdigit() and stripped[2:4] == ". ":
            document.add_paragraph(clean_inline_markdown(stripped[4:]), style="List Number")
            continue

        document.add_paragraph(clean_inline_markdown(line))

    document.save(docx_path)


if __name__ == "__main__":
    main()
